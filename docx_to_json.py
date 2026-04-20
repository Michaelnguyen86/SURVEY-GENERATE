"""
docx_to_json.py — DOCX Questionnaire → JSON Config Parser

Automatically converts a structured Word (.docx) questionnaire into
the survey_config.json format expected by survey_app.py.

Supports questionnaire structure:
  - Introductory paragraphs
  - Part I: Demographic questions (numbered paragraphs with option lists)
  - Parts II–IX: Likert-scale tables (7-column, 5-point scale)

Usage (command line):
    python docx_to_json.py "path/to/questionnaire.docx" [--output survey_config.json]

Usage (Python import):
    from docx_to_json import parse_docx_to_config
    cfg = parse_docx_to_config("questionnaire.docx", existing_meta)
"""

import argparse
import json
import os
import re
import sys

from docx import Document


# ─── KNOWN SECTION MAPPING ────────────────────────────────────────────────────
# Maps Vietnamese section keywords → construct IDs and variable prefixes
SECTION_KEYWORDS = [
    ("SEI",  "sei_",  "BẤT BÌNH ĐẲNG VỀ CẤU TRÚC VÀ KINH TẾ XÃ HỘI"),
    ("LWI",  "lwi_",  "BẤT BÌNH ĐẲNG VỀ GIỚI TÍNH VÀ LAO ĐỘNG"),
    ("PGI",  "pgi_",  "BẤT BÌNH ĐẲNG VỀ QUYỀN LỰC"),
    ("SCPO", "scpo_", "THỰC THI VÀ KẾT QUẢ CHUỖI CUNG ỨNG"),
    ("SESI", "sesi_", "TÁC ĐỘNG KINH TẾ"),
    ("HCSC", "hcsc_", "VĂN HÓA CHUỖI CUNG ỨNG"),
    ("GAI",  "gai_",  "GENERATIVE AI"),
    ("AAI",  "aai_",  "AGENTIC AI"),
    ("ESG",  "esg_",  "ESG"),
]

DEFAULT_SCALE_LABELS = [
    "Hoàn toàn không đồng ý",
    "Không đồng ý",
    "Trung lập",
    "Đồng ý",
    "Hoàn toàn đồng ý",
]


def _detect_section_id(title: str):
    """Return (id, prefix) by matching title to known keywords."""
    title_upper = title.upper()
    for sid, prefix, keyword in SECTION_KEYWORDS:
        if keyword in title_upper:
            return sid, prefix
    # Fallback: extract acronym in parentheses, e.g. "(SEI)"
    match = re.search(r"\(([A-Z]{2,6})\)", title)
    if match:
        acronym = match.group(1)
        return acronym, f"{acronym.lower()}_"
    return "SECT", "sect_"


def _parse_table_section(table, table_index: int) -> dict:
    """
    Parse a Likert table into a section dict.
    Expected table structure:
      Row 0-1 : merged title header
      Row 2   : scale label header (cols 2=low label, 6=high label)
      Row 3+  : questions (col 0=number, col 1=text, cols 2-6=scale 1-5)
    """
    # Get title from first non-empty cell in rows 0–1
    title = ""
    for row_idx in range(min(2, len(table.rows))):
        cell_text = table.cell(row_idx, 0).text.strip()
        if cell_text:
            title = cell_text
            break

    sec_id, prefix = _detect_section_id(title)

    questions = []
    for row in table.rows[3:]:
        cells = [c.text.strip() for c in row.cells]
        if not cells[0].isdigit():
            continue
        num  = int(cells[0])
        text = cells[1].strip()
        if text:
            questions.append({"number": num, "text": text})

    return {
        "id"              : sec_id,
        "title"           : f"PHẦN {table_index}: {title}",
        "short_title"     : f"{sec_id} — {title[:60]}",
        "description"     : title,
        "variable_prefix" : prefix,
        "scale"           : 5,
        "scale_labels"    : DEFAULT_SCALE_LABELS,
        "questions"       : questions,
    }


def _parse_demographic_paragraphs(paragraphs: list) -> dict:
    """
    Heuristically parse numbered demographic questions from plain paragraphs.
    Looks for lines matching "N. Question text" followed by option lines.
    """
    questions = []
    current_q = None
    current_opts = []

    # Regex: starts with a digit and period/dot
    q_pattern  = re.compile(r"^(\d+)\.\s+(.+)")
    # Common option separators: multiple spaces, tabs
    opt_split  = re.compile(r"\s{2,}|\t")

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        m = q_pattern.match(text)
        if m:
            # Save previous question
            if current_q is not None:
                questions.append(_build_demo_question(current_q, current_opts))
            current_q   = (int(m.group(1)), m.group(2).strip())
            current_opts = []
        elif current_q is not None:
            # Could be single option line or multi-option on one line
            parts = [p.strip() for p in opt_split.split(text) if p.strip()]
            current_opts.extend(parts)

    if current_q is not None:
        questions.append(_build_demo_question(current_q, current_opts))

    return {
        "title"    : "PHẦN I: THÔNG TIN VỀ ĐÁP VIÊN",
        "questions": questions,
    }


def _build_demo_question(q_tuple, opts: list) -> dict:
    num, text = q_tuple
    variable  = f"q_info_{num}"
    has_other = any("khác" in o.lower() or "other" in o.lower() for o in opts)

    # Build coding: option text → sequential integer
    coding = {opt: i + 1 for i, opt in enumerate(opts)}

    return {
        "number"  : num,
        "variable": variable,
        "text"    : text,
        "type"    : "radio" if opts else "text",
        "options" : opts,
        "coding"  : coding,
        "has_other": has_other,
    }


def _default_meta() -> dict:
    return {
        "title"         : "BẢNG CÂU HỎI KHẢO SÁT",
        "subtitle"      : "",
        "intro_text"    : "",
        "admin_password": "research2026",
        "version"       : "v1.0",
        "language"      : "vi",
    }


def parse_docx_to_config(docx_path: str, existing_meta: dict = None) -> dict:
    """
    Parse a DOCX questionnaire file and return a survey_config dict.

    Parameters
    ----------
    docx_path     : path to the .docx file
    existing_meta : if provided, keeps the existing survey_meta (title, password, etc.)

    Returns
    -------
    dict  survey_config compatible with survey_app.py
    """
    doc = Document(docx_path)

    meta = existing_meta.copy() if existing_meta else _default_meta()

    # Extract intro text from early paragraphs (before the first numbered question)
    intro_lines = []
    demo_paras  = []
    in_intro    = True
    q_pattern   = re.compile(r"^\d+\.")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if in_intro and not q_pattern.match(text):
            # Could be title, intro or instructions
            if len(text) > 10 and not text.isupper():
                intro_lines.append(text)
        else:
            in_intro = False
            demo_paras.append(para)

    if intro_lines and not meta.get("intro_text"):
        meta["intro_text"] = "\n\n".join(intro_lines[:4])

    # Parse demographics from paragraphs
    demo_section = _parse_demographic_paragraphs(demo_paras)

    # If parser found no demographic questions, add a placeholder
    if not demo_section["questions"]:
        demo_section["questions"] = [
            {
                "number": 1,
                "variable": "q_info_1",
                "text": "(Câu hỏi phân loại — vui lòng chỉnh sửa trong survey_config.json)",
                "type": "radio",
                "options": ["Lựa chọn 1", "Lựa chọn 2"],
                "coding": {"Lựa chọn 1": 1, "Lựa chọn 2": 2},
                "has_other": False,
            }
        ]

    # Parse Likert tables (skip table 0 which is usually the header/date table)
    likert_sections = []
    for t_idx, table in enumerate(doc.tables):
        if t_idx == 0:
            continue  # header table
        section = _parse_table_section(table, t_idx)
        if section["questions"]:
            likert_sections.append(section)

    return {
        "survey_meta"       : meta,
        "demographic_section": demo_section,
        "likert_sections"   : likert_sections,
    }


# ─── CLI ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Convert a DOCX questionnaire to survey_config.json"
    )
    parser.add_argument("docx_file", help="Path to the .docx questionnaire")
    parser.add_argument(
        "--output",
        default="survey_config_new.json",
        help="Output JSON file path (default: survey_config_new.json)",
    )
    args = parser.parse_args()

    if not os.path.exists(args.docx_file):
        print(f"ERROR: File not found: {args.docx_file}", file=sys.stderr)
        sys.exit(1)

    print(f"Parsing: {args.docx_file}")
    cfg = parse_docx_to_config(args.docx_file)

    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)

    n_demo    = len(cfg["demographic_section"]["questions"])
    n_likert  = sum(len(s["questions"]) for s in cfg["likert_sections"])
    n_sections = len(cfg["likert_sections"])

    print(f"Done! Output: {args.output}")
    print(f"  Demographic questions : {n_demo}")
    print(f"  Likert sections       : {n_sections}")
    print(f"  Likert items total    : {n_likert}")


if __name__ == "__main__":
    main()
